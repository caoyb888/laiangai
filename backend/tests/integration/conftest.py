"""
集成测试全局 Fixture
- 使用内存 SQLite + 真实 FastAPI app（不连外部服务）
- MinIO / Milvus 全部 Mock，见 CLAUDE.md §8.3
- 测试 JWT 由 create_access_token 生成，与真实 auth 逻辑一致
"""
import io
import pytest
import pytest_asyncio
from unittest.mock import AsyncMock, MagicMock, patch
from httpx import AsyncClient, ASGITransport
from docx import Document as DocxDocument

# SQLite 不支持 MySQL 专有类型，注册兼容映射（仅测试环境）
from sqlalchemy.dialects.sqlite.base import SQLiteTypeCompiler
if not hasattr(SQLiteTypeCompiler, "visit_LONGTEXT"):
    SQLiteTypeCompiler.visit_LONGTEXT = lambda self, type_, **kw: "TEXT"  # type: ignore[attr-defined]


# ── 生成测试用 JWT ────────────────────────────────────────────────────────────


def _make_test_token() -> str:
    from app.core.security import create_access_token
    return create_access_token(user_id="test-user-001", role="reviewer")


# ── Fixtures ──────────────────────────────────────────────────────────────────


@pytest_asyncio.fixture
async def client():
    """
    创建 httpx AsyncClient，使用 ASGI Transport 直连 FastAPI app。
    通过 patch 阻止 init_minio / init_milvus 连接外部服务。
    使用 SQLite 内存数据库并在测试前建表，测试后销毁。
    """
    with (
        patch("app.core.minio_client.init_minio", new_callable=AsyncMock),
        patch("app.core.milvus_client.init_milvus", new_callable=AsyncMock),
    ):
        from app.core.db import engine
        from app.models.base import Base
        # 建表（SQLite 测试库，每次测试周期重建）
        async with engine.begin() as conn:
            await conn.run_sync(Base.metadata.create_all)
        from app.main import app
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as c:
            yield c
        # 测试结束后清理表
        async with engine.begin() as conn:
            await conn.run_sync(Base.metadata.drop_all)


@pytest.fixture
def auth_headers() -> dict[str, str]:
    """返回带有效 Bearer Token 的请求头"""
    token = _make_test_token()
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def sample_docx() -> bytes:
    """生成最小合法 .docx 文件（脱敏虚构内容，见 CLAUDE.md §8.2）"""
    doc = DocxDocument()
    doc.add_heading("测试合同", level=1)
    doc.add_paragraph("第一条  本合同由[企业名称]甲方与乙方签订。")
    doc.add_paragraph("第二条  合同金额为[金额]，支付方式为月结。")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def mock_minio():
    """返回 MinIO Mock，put_object / bucket_exists 均成功"""
    m = MagicMock()
    m.bucket_exists.return_value = True
    m.put_object.return_value = None
    m.presigned_get_object.return_value = "http://minio.test/presigned-url"
    return m
