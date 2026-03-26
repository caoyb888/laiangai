import io
import re
import zipfile
from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.table import Table as DocxTable
from .base import BaseParser, ParsedDocument, DocumentMeta, TextBlock, TableBlock, BlockType
import structlog

logger = structlog.get_logger()

# 标题样式名映射
HEADING_STYLE_MAP: dict[str, int] = {
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
    "标题 1": 1, "标题 2": 2, "标题 3": 3,
    "标题1": 1, "标题2": 2, "标题3": 3,
    # 钢铁行业文档常见样式
    "一级标题": 1, "二级标题": 2, "三级标题": 3,
}


class DocxParser(BaseParser):

    async def parse(self, content: bytes, file_name: str) -> ParsedDocument:
        try:
            doc = DocxDocument(io.BytesIO(content))
        except Exception as e:
            logger.warning("python-docx 直接打开失败，尝试修复 Content_Types",
                           file_name=file_name, error=str(e))
            try:
                doc = self._open_with_content_type_fix(content, file_name)
                logger.info("Content_Types 修复成功", file_name=file_name)
            except Exception as e2:
                logger.warning("Content_Types 修复仍失败，改用 Tika 兜底",
                               file_name=file_name, error=str(e2))
                return self._parse_via_tika(content, file_name)

        blocks: list[TextBlock | TableBlock] = []

        index = 0
        section_stack: list[str] = []   # 维护章节路径

        # 遍历文档元素（段落 + 表格，保持顺序）
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = DocxParagraph(element, doc)
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""
                level = HEADING_STYLE_MAP.get(style_name, 0)

                if level > 0:
                    # 维护章节路径栈
                    section_stack = section_stack[:level - 1]
                    section_stack.append(text)
                    block_type = BlockType.HEADING
                elif style_name.lower() in ("list paragraph",):
                    block_type = BlockType.LIST_ITEM
                else:
                    block_type = BlockType.PARAGRAPH

                blocks.append(TextBlock(
                    block_type=block_type,
                    text=text,
                    level=level,
                    index=index,
                    section_path="/".join(section_stack[:-1]) if level > 0 else "/".join(section_stack),
                    style_name=style_name,
                ))
                index += 1

            elif tag == "tbl":
                table = DocxTable(element, doc)
                rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows.append(row_data)
                blocks.append(TableBlock(
                    rows=rows,
                    index=index,
                    section_path="/".join(section_stack),
                ))
                index += 1

        # 提取元数据
        core_props = doc.core_properties
        word_count = sum(len(b.text) for b in blocks if isinstance(b, TextBlock))
        title = core_props.title or (blocks[0].text if blocks else "")

        meta = DocumentMeta(
            file_name=file_name,
            file_type="docx",
            page_count=None,     # docx 不直接提供页数
            word_count=word_count,
            title=title[:512],
            author=core_props.author or "",
        )

        raw_text = self._build_raw_text(blocks)

        logger.info("Word 解析完成",
                    file_name=file_name,
                    blocks=len(blocks),
                    word_count=word_count)

        return ParsedDocument(meta=meta, blocks=blocks, raw_text=raw_text)

    def _open_with_content_type_fix(self, content: bytes, file_name: str) -> DocxDocument:
        """
        修复 [Content_Types].xml 后重新尝试打开（针对 WPS/模板另存的非标 OOXML）。
        若仍失败则抛出，由上层调用 _parse_via_tika。
        """
        DOCUMENT_CT = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document.main+xml"
        )
        buf_in = io.BytesIO(content)
        buf_out = io.BytesIO()
        with zipfile.ZipFile(buf_in, "r") as zin, \
             zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    ct_text = data.decode("utf-8")
                    if "word/document.xml" not in ct_text:
                        override = (
                            f'<Override PartName="/word/document.xml" '
                            f'ContentType="{DOCUMENT_CT}"/>'
                        )
                        ct_text = ct_text.replace("</Types>", override + "</Types>")
                        data = ct_text.encode("utf-8")
                zout.writestr(item, data)
        buf_out.seek(0)
        return DocxDocument(buf_out)

    def _parse_via_tika(self, content: bytes, file_name: str) -> ParsedDocument:
        """
        Tika 兜底解析（处理老格式 .doc / 非标 OOXML）。
        Tika 返回纯文本，按段落转为 TextBlock 列表。
        """
        from tika import parser as tika_parser
        parsed = tika_parser.from_buffer(content)
        text: str = parsed.get("content") or ""
        if not text.strip():
            raise ValueError("Tika 解析结果为空，文件可能不含可提取的文本内容")

        logger.info("Tika 解析成功", file_name=file_name, text_len=len(text))

        # 将纯文本按段落切分为 TextBlock
        _heading_re = re.compile(
            r"^(第[一二三四五六七八九十百\d]+[条款章节项]"
            r"|[一二三四五六七八九十]+[、．.]\s*\S"
            r"|\d+[、．.]\s*\S)"
        )
        blocks: list[TextBlock] = []
        index = 0
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            is_heading = bool(_heading_re.match(line))
            blocks.append(TextBlock(
                block_type=BlockType.HEADING if is_heading else BlockType.PARAGRAPH,
                text=line,
                level=1 if is_heading else 0,
                index=index,
                section_path="",
                style_name="",
            ))
            index += 1

        word_count = sum(len(b.text) for b in blocks)
        meta = DocumentMeta(
            file_name=file_name,
            file_type="doc",
            page_count=None,
            word_count=word_count,
            title=blocks[0].text[:128] if blocks else file_name,
            author="",
        )
        raw_text = "\n".join(b.text for b in blocks)
        return ParsedDocument(meta=meta, blocks=blocks, raw_text=raw_text)
