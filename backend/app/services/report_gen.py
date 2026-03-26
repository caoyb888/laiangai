"""
比对报告生成：支持导出 Word(.docx) 和 PDF 格式
报告格式符合集团文件规范，见方案 §4.2.5
"""
import io
import os
from datetime import datetime
from xml.sax.saxutils import escape as _xml_escape

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import structlog

logger = structlog.get_logger()


class ReportGenerator:

    TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "report_templates")

    # 差异等级颜色（见 CLAUDE.md §6.2 颜色体系）
    LEVEL_COLORS = {
        "CRITICAL": RGBColor(0xD3, 0x2F, 0x2F),
        "MAJOR":    RGBColor(0xF5, 0x7C, 0x00),
        "MINOR":    RGBColor(0x38, 0x8E, 0x3C),
    }

    async def generate_docx(self, report_data: dict) -> bytes:
        """生成 Word 格式报告"""
        doc = DocxDocument()

        # ── 封面 ──────────────────────────────────────────────────────────
        title_para = doc.add_heading("文档比对报告", level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"任务名称：{report_data.get('task_name', '未命名')}")
        doc.add_paragraph(f"文档A：{report_data.get('doc_a_name', '')}")
        doc.add_paragraph(f"文档B：{report_data.get('doc_b_name', '')}")
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_page_break()

        # ── 摘要 ──────────────────────────────────────────────────────────
        doc.add_heading("一、比对摘要", level=1)
        summary = report_data.get("summary", {})
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = "Table Grid"
        rows_data = [
            ("总差异数", str(summary.get("total_diffs", 0))),
            ("重大差异（CRITICAL）", str(summary.get("critical_diffs", 0))),
            ("一般差异（MAJOR）", str(summary.get("major_diffs", 0))),
            ("格式差异（MINOR）", str(summary.get("minor_diffs", 0))),
        ]
        for i, (label, value) in enumerate(rows_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value

        doc.add_page_break()

        # ── 差异详情 ──────────────────────────────────────────────────────
        doc.add_heading("二、差异详情", level=1)
        for item in report_data.get("diff_items", []):
            level = item.get("diff_level", "MINOR")
            color = self.LEVEL_COLORS.get(level, RGBColor(0, 0, 0))

            # 差异标题（带颜色）
            p = doc.add_paragraph()
            run = p.add_run(f"[{level}] 差异 #{item.get('seq_no', 0) + 1}")
            run.bold = True
            run.font.color.rgb = color

            if item.get("doc_a_text"):
                doc.add_paragraph(f"原文：{item['doc_a_text'][:500]}")
            if item.get("doc_b_text"):
                doc.add_paragraph(f"修改后：{item['doc_b_text'][:500]}")
            if item.get("semantic_desc"):
                p = doc.add_paragraph()
                p.add_run("分析：").bold = True
                p.add_run(item["semantic_desc"])
            if item.get("risk_keywords"):
                p = doc.add_paragraph()
                p.add_run("风险关键词：").bold = True
                p.add_run(item["risk_keywords"])
            doc.add_paragraph("─" * 50)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def generate_pdf(self, report_data: dict) -> bytes:
        """生成 PDF 格式报告（使用 ReportLab）"""
        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=72, rightMargin=72, topMargin=72, bottomMargin=72
        )

        # 注册中文字体：按优先级查找可用字体
        _font_candidates = [
            ("/host-fonts/wqy-zenhei.ttc", "WQYZenHei"),
            ("/host-fonts/wqy-microhei.ttc", "WQYMicroHei"),
            ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "WQYZenHei"),
            ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", "WQYMicroHei"),
            (os.path.join(self.TEMPLATE_DIR, "fonts", "SimSun.ttf"), "SimSun"),
        ]
        base_font = "Helvetica"
        for _fpath, _fname in _font_candidates:
            if os.path.exists(_fpath):
                try:
                    from reportlab.pdfbase.ttfonts import TTFont as _TTFont
                    pdfmetrics.registerFont(_TTFont(_fname, _fpath, subfontIndex=0))
                    base_font = _fname
                    logger.info("PDF 字体加载成功", font=_fname, path=_fpath)
                    break
                except Exception as _e:
                    logger.warning("PDF 字体加载失败，尝试下一个", font=_fname, error=str(_e))

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "TitleCN", parent=styles["Title"],
            fontName=base_font, fontSize=18, leading=26, spaceAfter=6
        )
        h1_style = ParagraphStyle(
            "H1CN", parent=styles["Heading1"],
            fontName=base_font, fontSize=14, leading=20, spaceBefore=16, spaceAfter=8
        )
        normal_style = ParagraphStyle(
            "NormalCN", parent=styles["Normal"],
            fontName=base_font, fontSize=10, leading=16
        )
        bold_style = ParagraphStyle(
            "BoldCN", parent=styles["Normal"],
            fontName=base_font, fontSize=10, leading=16, fontWeight="bold"
        )
        label_style = ParagraphStyle(
            "LabelCN", parent=styles["Normal"],
            fontName=base_font, fontSize=9, leading=14, textColor=colors.HexColor("#555555")
        )

        # 差异等级颜色
        _level_colors = {
            "CRITICAL": colors.HexColor("#D32F2F"),
            "MAJOR":    colors.HexColor("#F57C00"),
            "MINOR":    colors.HexColor("#388E3C"),
        }

        story = []

        # ── 封面 ──────────────────────────────────────────────────────────
        story.append(Spacer(1, 40))
        story.append(Paragraph("莱钢集团 文档比对报告", title_style))
        story.append(Spacer(1, 16))
        meta_data = [
            ["任务名称", report_data.get("task_name", "未命名")],
            ["文档A（基准）", report_data.get("doc_a_name", "")],
            ["文档B（对比）", report_data.get("doc_b_name", "")],
            ["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ]
        meta_table = Table(meta_data, colWidths=[100, 340])
        meta_table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), base_font),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#555555")),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 24))

        # ── 比对摘要 ──────────────────────────────────────────────────────
        story.append(Paragraph("一、比对摘要", h1_style))
        summary = report_data.get("summary", {})
        table_data = [
            ["项目", "数值"],
            ["总差异数", str(summary.get("total_diffs", 0))],
            ["重大差异（CRITICAL）", str(summary.get("critical_diffs", 0))],
            ["一般差异（MAJOR）", str(summary.get("major_diffs", 0))],
            ["格式差异（MINOR）", str(summary.get("minor_diffs", 0))],
        ]
        t = Table(table_data, colWidths=[300, 140])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#455A64")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("FONTNAME", (0, 0), (-1, -1), base_font),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F5F5")]),
        ]))
        story.append(t)
        story.append(Spacer(1, 24))

        # ── 差异详情 ──────────────────────────────────────────────────────
        story.append(Paragraph("二、差异详情", h1_style))

        from reportlab.platypus import HRFlowable
        for item in report_data.get("diff_items", []):
            level = item.get("diff_level", "MINOR")
            level_color = _level_colors.get(level, colors.black)
            seq = item.get("seq_no", 0) + 1

            # 差异标题行
            level_labels = {"CRITICAL": "重大差异", "MAJOR": "一般差异", "MINOR": "格式差异"}
            header_style = ParagraphStyle(
                f"DiffHeader{seq}", parent=styles["Normal"],
                fontName=base_font, fontSize=10, leading=16,
                textColor=level_color, spaceBefore=10
            )
            story.append(Paragraph(
                f"【{level_labels.get(level, level)}】差异 #{seq}",
                header_style
            ))

            if item.get("doc_a_text"):
                story.append(Paragraph("原文：", label_style))
                story.append(Paragraph(_xml_escape(item["doc_a_text"][:500]), normal_style))
            if item.get("doc_b_text"):
                story.append(Paragraph("修改后：", label_style))
                story.append(Paragraph(_xml_escape(item["doc_b_text"][:500]), normal_style))
            if item.get("semantic_desc"):
                story.append(Paragraph("分析：", label_style))
                story.append(Paragraph(_xml_escape(item["semantic_desc"]), normal_style))
            if item.get("risk_keywords"):
                story.append(Paragraph(f"风险关键词：{_xml_escape(item['risk_keywords'])}", label_style))

            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.HexColor("#DDDDDD"), spaceAfter=4))

        doc.build(story)
        return buf.getvalue()
